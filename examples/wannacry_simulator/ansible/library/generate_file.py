#!/usr/bin/python
# -*- coding: utf-8 -*-

from __future__ import absolute_import, division, print_function
__metaclass__ = type

DOCUMENTATION = r'''
---
module: generate_file

short_description: Generates files of different types with custom content and metadata

description:
  - Generic module to generate files of multiple formats
  - Supports: PDF, images, DOCX, audio, video, compressed files, plain text, executables
  - Allows specifying explicit content or auto-generating it
  - Controls timestamps (crtime, mtime, atime) and permissions

options:
  name:
    description: Name of the file to generate
    required: true
    type: str
  
  extension:
    description: Extension/MIME type of the file
    required: true
    type: str
    choices: [pdf, jpg, png, docx, doc, odt, mp3, wav, mp4, avi, zip, tar.gz, txt, sh, py]
  
  path:
    description: Full path where to save the file
    required: true
    type: str
  
  content:
    description: Explicit file content (string or base64)
    required: false
    type: str
  
  size:
    description: Target size in bytes (alternative to content)
    required: false
    type: int
  
  template:
    description: Content template (txt, img, code, etc.)
    required: false
    type: str
    choices: [text, image, code, document, audio, video]
  
  creation_date:
    description: Creation timestamp (crtime)
    required: false
    type: str
  
  modification_date:
    description: Modification timestamp (mtime)
    required: false
    type: str
  
  access_date:
    description: Access timestamp (atime)
    required: false
    type: str
  
  permissions:
    description: File permissions (octal format)
    required: false
    type: str
    default: "0644"
  
  owner:
    description: File owner
    required: false
    type: str
  
  group:
    description: File group
    required: false
    type: str
'''

import os
import sys
import json
import random
import subprocess
from datetime import datetime
from ansible.module_utils.basic import AnsibleModule


def generate_pdf(path, content=None, size=None):
    """Generates a PDF file"""
    try:
        import tempfile
        from faker import Faker
        fake = Faker('es_ES')
        
        # Markdown content
        if content:
            markdown = content
        else:
            markdown = f"""---
title: "{fake.sentence()}"
author: "{fake.name()}"
date: "{datetime.now().strftime('%Y-%m-%d')}"
---

# {fake.sentence()}

{fake.paragraph()}

## Section 1

{fake.text(max_nb_chars=200)}
"""
        
        # Create temp file
        with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False, encoding='utf-8') as f:
            f.write(markdown)
            temp_md = f.name
        
        # Convert with pandoc
        result = subprocess.run([
            'pandoc', temp_md, '-o', path,
            '--pdf-engine=pdflatex'
        ], capture_output=True, text=True)
        
        os.unlink(temp_md)
        
        if result.returncode != 0:
            # Fallback: minimal PDF file
            with open(path, 'wb') as f:
                f.write(b'%PDF-1.4\n%\xE2\xE3\xCF\xD3\n')
        
        return True, os.path.getsize(path)
    except Exception as e:
        return False, str(e)


def generate_image(path, extension, size=None):
    """Generates an image (JPG, PNG)"""
    try:
        from PIL import Image, ImageDraw, ImageFont
        import random
        
        width = random.randint(400, 1920)
        height = random.randint(300, 1080)
        
        # Create image with random color
        color = (random.randint(0, 255), random.randint(0, 255), random.randint(0, 255))
        img = Image.new('RGB', (width, height), color)
        
        # Add text
        draw = ImageDraw.Draw(img)
        text = f"Generated {datetime.now()}"
        draw.text((50, 50), text, fill=(255, 255, 255))
        
        # Save
        if extension == 'png':
            img.save(path, 'PNG')
        else:
            img.save(path, 'JPEG', quality=85)
        
        return True, os.path.getsize(path)
    except Exception as e:
        # Fallback: minimal image
        with open(path, 'wb') as f:
            if extension == 'png':
                f.write(b'\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01\x08\x02\x00\x00\x00\x90wS\xde')
            else:
                f.write(b'\xff\xd8\xff\xe0\x00\x10JFIF')
        return True, os.path.getsize(path)


def generate_text(path, content=None, size=None):
    """Generates plain text file"""
    try:
        from faker import Faker
        fake = Faker('es_ES')
        
        if content:
            text = content
        else:
            text = "\n\n".join([fake.paragraph() for _ in range(10)])
        
        with open(path, 'w', encoding='utf-8') as f:
            f.write(text)
        
        return True, os.path.getsize(path)
    except Exception as e:
        return False, str(e)


def generate_docx(path, content=None):
    """Generates DOCX document"""
    try:
        from docx import Document
        from faker import Faker
        fake = Faker('es_ES')
        
        doc = Document()
        doc.add_heading(fake.sentence(), 0)
        
        for _ in range(5):
            doc.add_paragraph(fake.paragraph())
        
        doc.save(path)
        return True, os.path.getsize(path)
    except:
        # Fallback: create ZIP with minimal DOCX structure
        import zipfile
        with zipfile.ZipFile(path, 'w') as zf:
            zf.writestr('[Content_Types].xml', '<?xml version="1.0"?><Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types"></Types>')
        return True, os.path.getsize(path)


def generate_compressed(path, extension):
    """Generates compressed file (ZIP, TAR.GZ)"""
    try:
        import zipfile
        import tarfile
        from faker import Faker
        fake = Faker('es_ES')
        
        if extension == 'zip':
            with zipfile.ZipFile(path, 'w') as zf:
                for i in range(3):
                    zf.writestr(f'file_{i}.txt', fake.text())
        elif extension in ['tar.gz', 'tar']:
            with tarfile.open(path, 'w:gz' if extension == 'tar.gz' else 'w') as tf:
                import tempfile
                for i in range(3):
                    with tempfile.NamedTemporaryFile(mode='w', delete=False) as tmp:
                        tmp.write(fake.text())
                        tmp.flush()
                        tf.add(tmp.name, arcname=f'file_{i}.txt')
                        os.unlink(tmp.name)
        
        return True, os.path.getsize(path)
    except Exception as e:
        return False, str(e)


def generate_executable(path, extension):
    """Generates executable script"""
    try:
        if extension == 'sh':
            content = """#!/bin/bash
# Automatically generated script
echo "Executable script generated by Ansible"
date
"""
        elif extension == 'py':
            content = """#!/usr/bin/env python3
# Automatically generated script
import sys
print("Python script generated by Ansible")
sys.exit(0)
"""
        else:
            content = "#!/bin/sh\necho 'Generated script'\n"
        
        with open(path, 'w') as f:
            f.write(content)
        
        os.chmod(path, 0o755)
        return True, os.path.getsize(path)
    except Exception as e:
        return False, str(e)


def apply_timestamps(path, mtime=None, atime=None):
    """Applies timestamps to file"""
    import subprocess
    from datetime import datetime
    
    # Convert strings to timestamps
    def parse_date(date_str):
        if not date_str:
            return None
        try:
            return datetime.strptime(date_str, "%Y-%m-%d %H:%M:%S").timestamp()
        except:
            return None
    
    mtime_ts = parse_date(mtime)
    atime_ts = parse_date(atime)
    
    # Apply mtime and atime with utime
    if mtime_ts or atime_ts:
        atime_val = atime_ts if atime_ts else os.stat(path).st_atime
        mtime_val = mtime_ts if mtime_ts else os.stat(path).st_mtime
        os.utime(path, (atime_val, mtime_val))
    
    # crtime can only be modified with debugfs (requires root)
    # For now we omit it, but could be implemented
    
    return True


def main():
    module = AnsibleModule(
        argument_spec=dict(
            name=dict(type='str', required=True),
            extension=dict(type='str', required=True),
            path=dict(type='str', required=True),
            content=dict(type='str', required=False, default=''),
            size=dict(type='int', required=False, default=0),
            template=dict(type='str', required=False, default=''),
            creation_date=dict(type='str', required=False, default=''),
            modification_date=dict(type='str', required=False, default=''),
            access_date=dict(type='str', required=False, default=''),
            permissions=dict(type='str', required=False, default='0644'),
            owner=dict(type='str', required=False, default=''),
            group=dict(type='str', required=False, default='')
        ),
        supports_check_mode=True
    )
    
    name = module.params['name']
    extension = module.params['extension']
    path = module.params['path']
    content = module.params['content']
    size = module.params['size']
    
    # Create directory if it doesn't exist
    os.makedirs(os.path.dirname(path), exist_ok=True)
    
    # Generate according to type
    success = False
    file_size = 0
    
    if extension in ['pdf']:
        success, file_size = generate_pdf(path, content, size)
    elif extension in ['jpg', 'jpeg', 'png']:
        success, file_size = generate_image(path, extension, size)
    elif extension in ['txt', 'log', 'csv']:
        success, file_size = generate_text(path, content, size)
    elif extension in ['docx', 'doc', 'odt']:
        success, file_size = generate_docx(path, content)
    elif extension in ['zip', 'tar.gz', 'tar']:
        success, file_size = generate_compressed(path, extension)
    elif extension in ['sh', 'py', 'bash']:
        success, file_size = generate_executable(path, extension)
    else:
        # Generic file
        with open(path, 'wb') as f:
            f.write(b'Generic file content\n')
        success = True
        file_size = os.path.getsize(path)
    
    if not success:
        module.fail_json(msg=f"Error generating file: {file_size}")
    
    # Apply timestamps
    apply_timestamps(
        path,
        # module.params['creation_date'],
        module.params['modification_date'],
        module.params['access_date']
    )
    
    # Apply permissions
    if module.params['permissions']:
        os.chmod(path, int(module.params['permissions'], 8))
    
    # Change owner/group
    if module.params['owner'] or module.params['group']:
        import pwd, grp
        uid = pwd.getpwnam(module.params['owner']).pw_uid if module.params['owner'] else -1
        gid = grp.getgrnam(module.params['group']).gr_gid if module.params['group'] else -1
        if uid != -1 or gid != -1:
            os.chown(path, uid, gid)
    
    module.exit_json(
        changed=True,
        path=path,
        name=name,
        extension=extension,
        size_bytes=file_size,
        size_kb=round(file_size / 1024, 2)
    )


if __name__ == '__main__':
    main()
