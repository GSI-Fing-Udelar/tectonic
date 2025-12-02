#!/usr/bin/python
# -*- coding: utf-8 -*-

from __future__ import absolute_import, division, print_function
__metaclass__ = type

DOCUMENTATION = r'''
---
module: generate_files_bulk

short_description: Bulk file generation orchestrator

description:
  - Generates multiple files using layer 1 primitives
  - Total file count control
  - Distribution by types (percentages)
  - Directory structure (flat or tree)
  - Name pattern with templates

options:
  count:
    description: Total number of files to generate
    required: true
    type: int
  
  distribution:
    description: Type distribution with percentages
    required: false
    type: dict
    default: {pdf: 30, docx: 40, txt: 30}
  
  structure:
    description: Directory structure (flat or tree)
    required: false
    type: str
    choices: [flat, tree]
    default: flat
  
  base_directory:
    description: Root directory where files will be generated
    required: true
    type: str
  
  name_pattern:
    description: Template for names (supports {n}, {date}, {file_type}, {faker_*})
    required: false
    type: str
    default: "document_{n}"
    notes:
      - "Supported variables: {n}, {date}, {time}, {file_type}"
      - "Faker variables: {faker_name}, {faker_company}, {faker_city}, {faker_word}, {faker_job}"
      - "Each file gets UNIQUE faker values using seed + file_index"
  
  tree_depth:
    description: Maximum depth of directory tree
    required: false
    type: int
    default: 3
  
  subdirs_per_level:
    description: Number of subdirectories per level in tree
    required: false
    type: int
    default: 3
  
  deleted_ratio:
    description: Percentage of files to generate as deleted (state:absent)
    required: false
    type: int
    default: 0
  
  faker_seed:
    description: Seed for Faker random generation (for reproducibility)
    required: false
    type: int
    default: 1
'''

import os
import sys
import json
import random
import subprocess
from datetime import datetime
from ansible.module_utils.basic import AnsibleModule

import tempfile
# Try to import optional libraries
try:
    from openpyxl import Workbook
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False



def distribute_files(count, distribution):
    """
    Distributes total count according to percentages
                        # Temporary DEBUG - use a safe temp file per uid and handle write permission errors
                        DEBUG_LOG = os.path.join(tempfile.gettempdir(), f"xlsx_debug_{os.geteuid()}.log")
                        def _append_debug_line(text):
                            try:
                                with open(DEBUG_LOG, 'a') as debug_f:
                                    debug_f.write(text)
                            except PermissionError:
                                # Avoid failing the module; optionally we could call module.warn but module may be undefined in some contexts
                                try:
                                    module.warn(f"Cannot write XLSX debug log: {DEBUG_LOG}")
                                except Exception:
                                    pass
                        _append_debug_line(f"=== XLSX #{global_index} ===\n")
                        _append_debug_line(f"OPENPYXL_AVAILABLE: {OPENPYXL_AVAILABLE}\n")
                        _append_debug_line(f"full_path: {full_path}\n")
    Returns:
        dict with {type: count}
    """
    result = {}
    total_percentage = sum(distribution.values())
    
    # Normalize percentages
    distribution_norm = {k: v / total_percentage for k, v in distribution.items()}
    
    # Calculate counts
    remaining = count
    for file_type, percentage in list(distribution_norm.items())[:-1]:
        qty = int(count * percentage)
        result[file_type] = qty
        remaining -= qty
    
    # Last type receives the rest
    last_type = list(distribution_norm.keys())[-1]
    result[last_type] = remaining
    
    return result


def generate_flat_structure(base_directory, count):
    """
    Generates flat structure (all in base_directory)
    
    Returns:
        list of paths where files will be created
    """
    os.makedirs(base_directory, exist_ok=True)
    return [base_directory] * count


def generate_tree_structure(base_directory, count, depth=3, subdirs_per_level=3, seed=1):
    """
    Generates tree structure with random subdirectories
    
    Returns:
        list of paths where files will be created
    """
    from faker import Faker
    Faker.seed(seed)
    fake = Faker('en_US')
    
    # Create directories
    directories = [base_directory]
    
    for level in range(depth):
        new_dirs = []
        for parent_dir in directories[-subdirs_per_level:]:
            for i in range(random.randint(1, subdirs_per_level)):
                dir_name = fake.word()
                new_dir = os.path.join(parent_dir, f"{dir_name}_{i}")
                os.makedirs(new_dir, exist_ok=True)
                new_dirs.append(new_dir)
        directories.extend(new_dirs)
    
    # Distribute files randomly among directories
    paths = []
    for _ in range(count):
        paths.append(random.choice(directories))
    
    return paths


def expand_pattern(pattern, index, file_type, faker_enabled=True, seed=1):
    """
    Expands name pattern with variables
    
    Supported variables:
        {n} - file number
        {date} - current date
        {time} - current time
        {file_type} - file type
        {faker_name} - random name
        {faker_company} - random company
        {faker_city} - random city
        {faker_word} - random word
        {faker_job} - random job title
    
    IMPORTANT: Uses seed + index to generate DIFFERENT values per file
    while maintaining reproducibility with the same seed.
    
    Args:
        pattern: Name pattern template
        index: File index (must be >= 0)
        file_type: File extension
        faker_enabled: Enable faker variables
        seed: Base seed for reproducibility
    """
    from datetime import datetime
    
    result = pattern
    
    # Basic variables
    result = result.replace('{n}', str(index))
    result = result.replace('{date}', datetime.now().strftime('%Y%m%d'))
    result = result.replace('{time}', datetime.now().strftime('%H%M%S'))
    result = result.replace('{file_type}', file_type)
    
    # Faker variables - UNIQUE per file using seed + index
    if faker_enabled and '{faker_' in result:
        try:
            from faker import Faker
            # Use seed + index to get different values per file
            # but still reproducible with same seed
            Faker.seed(seed + index)
            fake = Faker('en_US')
            
            # Helper function to sanitize faker output (remove /, \, :, etc)
            def sanitize(text):
                """Remove characters that are invalid in filenames"""
                invalid_chars = ['/', '\\', ':', '*', '?', '"', '<', '>', '|']
                for char in invalid_chars:
                    text = text.replace(char, '_')
                return text.replace(' ', '_')
            
            if '{faker_name}' in result:
                result = result.replace('{faker_name}', sanitize(fake.name()))
            if '{faker_company}' in result:
                result = result.replace('{faker_company}', sanitize(fake.company()))
            if '{faker_city}' in result:
                result = result.replace('{faker_city}', sanitize(fake.city()))
            if '{faker_word}' in result:
                result = result.replace('{faker_word}', sanitize(fake.word()))
            if '{faker_job}' in result:
                result = result.replace('{faker_job}', sanitize(fake.job()))
        except:
            pass
    
    return result


def main():
    module = AnsibleModule(
        argument_spec=dict(
            count=dict(type='int', required=True),
            distribution=dict(type='dict', required=False, default={'pdf': 30, 'docx': 40, 'txt': 30}),
            structure=dict(type='str', required=False, default='flat', choices=['flat', 'tree']),
            base_directory=dict(type='str', required=True),
            name_pattern=dict(type='str', required=False, default='document_{n}'),
            tree_depth=dict(type='int', required=False, default=3),
            subdirs_per_level=dict(type='int', required=False, default=3),
            default_permissions=dict(type='str', required=False, default='0644'),
            deleted_ratio=dict(type='int', required=False, default=0),
            faker_seed=dict(type='int', required=False, default=1)
        ),
        supports_check_mode=True
    )
    
    count = module.params['count']
    distribution = module.params['distribution']
    structure = module.params['structure']
    base_directory = module.params['base_directory']
    name_pattern = module.params['name_pattern']
    deleted_ratio = module.params['deleted_ratio']
    faker_seed = module.params['faker_seed']
    
    # Set global seeds for reproducibility
    random.seed(faker_seed)
    try:
        from faker import Faker
        Faker.seed(faker_seed)
    except ImportError:
        pass
    
    # Validations
    if count <= 0:
        module.fail_json(msg="Count must be greater than 0")
    
    if not distribution:
        module.fail_json(msg="Distribution cannot be empty")
    
    # 1. Calculate file distribution by type
    counts_per_type = distribute_files(count, distribution)
    
    # 2. Generate directory structure
    if structure == 'flat':
        paths = generate_flat_structure(base_directory, count)
    else:
        paths = generate_tree_structure(
            base_directory,
            count,
            module.params['tree_depth'],
            module.params['subdirs_per_level'],
            faker_seed
        )
    
    # 3. Generate files
    generated_files = []
    global_index = 1
    
    for file_type, type_count in counts_per_type.items():
        for i in range(type_count):
            # Expand name pattern
            name = expand_pattern(name_pattern, global_index, file_type, True, faker_seed)
            full_name = f"{name}.{file_type}"
            
            # Select path
            path_dir = paths[global_index - 1]
            full_path = os.path.join(path_dir, full_name)
            
            # Create file according to type
            try:
                success = False
                size = 0
                
                if file_type == 'pdf':
                    # PDF with pandoc
                    import tempfile
                    from faker import Faker
                    Faker.seed(faker_seed)
                    fake = Faker('en_US')
                    markdown = f"# {fake.sentence()}\n\n{fake.paragraph()}"
                    with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False, encoding='utf-8') as f:
                        f.write(markdown)
                        temp_md = f.name
                    
                    result = subprocess.run(['pandoc', temp_md, '-o', full_path, '--pdf-engine=pdflatex'], 
                                          capture_output=True, text=True, timeout=10)
                    os.unlink(temp_md)
                    
                    if result.returncode != 0:
                        # Fallback: minimal valid PDF
                        with open(full_path, 'wb') as f:
                            f.write(b'%PDF-1.4\n%\xE2\xE3\xCF\xD3\n1 0 obj\n<</Type/Catalog/Pages 2 0 R>>\nendobj\n')
                    size = os.path.getsize(full_path)
                    success = True
                    
                elif file_type in ['jpg', 'jpeg', 'png']:
                    # Image with PIL
                    from PIL import Image, ImageDraw
                    width, height = random.randint(400, 800), random.randint(300, 600)
                    color = (random.randint(0, 255), random.randint(0, 255), random.randint(0, 255))
                    img = Image.new('RGB', (width, height), color)
                    draw = ImageDraw.Draw(img)
                    draw.text((50, 50), f"File #{global_index}", fill=(255, 255, 255))
                    
                    if file_type == 'png':
                        img.save(full_path, 'PNG')
                    else:
                        img.save(full_path, 'JPEG', quality=85)
                    size = os.path.getsize(full_path)
                    success = True
                    
                elif file_type in ['txt', 'log', 'csv']:
                    # Plain text
                    from faker import Faker
                    Faker.seed(faker_seed)
                    fake = Faker('en_US')
                    text = "\n\n".join([fake.paragraph() for _ in range(3)])
                    with open(full_path, 'w', encoding='utf-8') as f:
                        f.write(text)
                    size = os.path.getsize(full_path)
                    success = True
                    
                elif file_type in ['docx', 'doc', 'odt']:
                    # DOCX
                    try:
                        from docx import Document
                        from faker import Faker
                        Faker.seed(faker_seed)
                        fake = Faker('en_US')
                        doc = Document()
                        doc.add_heading(fake.sentence(), 0)
                        for _ in range(3):
                            doc.add_paragraph(fake.paragraph())
                        doc.save(full_path)
                        size = os.path.getsize(full_path)
                        success = True
                    except:
                        # Fallback: ZIP with minimal DOCX structure
                        import zipfile
                        with zipfile.ZipFile(full_path, 'w') as zf:
                            zf.writestr('[Content_Types].xml', '<?xml version="1.0"?><Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types"></Types>')
                        size = os.path.getsize(full_path)
                        success = True
                
                elif file_type in ['xlsx', 'xls']:
                    # XLSX - Excel
                    xlsx_created = False
                    
                    # Temporary DEBUG - use a safe temp file per uid and handle write permission errors
                    DEBUG_LOG = os.path.join(tempfile.gettempdir(), f"xlsx_debug_{os.geteuid()}.log")
                    def _append_debug_line(text):
                        try:
                            with open(DEBUG_LOG, 'a') as debug_f:
                                debug_f.write(text)
                        except PermissionError:
                            try:
                                module.warn(f"Cannot write XLSX debug log: {DEBUG_LOG}")
                            except Exception:
                                pass
                    _append_debug_line(f"=== XLSX #{global_index} ===\n")
                    _append_debug_line(f"OPENPYXL_AVAILABLE: {OPENPYXL_AVAILABLE}\n")
                    _append_debug_line(f"full_path: {full_path}\n")
                    
                    if OPENPYXL_AVAILABLE:
                        try:
                            from faker import Faker
                            Faker.seed(faker_seed)
                            fake = Faker('en_US')
                            
                            wb = Workbook()
                            ws = wb.active
                            ws.title = "Sheet1"
                            
                            # Headers
                            headers = ['ID', 'Name', 'Email', 'City', 'Date']
                            ws.append(headers)
                            
                            # Data rows (10)
                            for i in range(10):
                                ws.append([
                                    i + 1,
                                    fake.name(),
                                    fake.email(),
                                    fake.city(),
                                    str(fake.date())
                                ])
                            
                            wb.save(full_path)
                            size = os.path.getsize(full_path)
                            success = True
                            xlsx_created = True
                            
                            # DEBUG
                            _append_debug_line(f"✓ openpyxl SUCCESS, size={size}\n\n")
                                
                        except Exception as e:
                            # Try fallback
                            xlsx_created = False
                            
                            # DEBUG
                            _append_debug_line(f"✗ openpyxl FAILED: {type(e).__name__}: {e}\n\n")
                    
                    if not xlsx_created:
                        # Fallback: CSV with xlsx extension
                        from faker import Faker
                        Faker.seed(faker_seed)
                        fake = Faker('en_US')
                        
                        csv_content = "ID,Name,Email,City,Date\n"
                        for i in range(10):
                            csv_content += f"{i+1},{fake.name()},{fake.email()},{fake.city()},{fake.date()}\n"
                        
                        with open(full_path, 'w', encoding='utf-8') as f:
                            f.write(csv_content)
                        
                        size = os.path.getsize(full_path)
                        success = True
                        
                        # DEBUG
                        _append_debug_line(f"\u2192 FALLBACK CSV used, size={size}\n\n")
                    
                elif file_type in ['zip', 'tar.gz', 'tar']:
                    # Compressed files
                    import zipfile
                    import tarfile
                    from faker import Faker
                    Faker.seed(faker_seed)
                    fake = Faker('en_US')
                    
                    if file_type == 'zip':
                        with zipfile.ZipFile(full_path, 'w') as zf:
                            for i in range(3):
                                zf.writestr(f'file_{i}.txt', fake.text())
                    else:
                        with tarfile.open(full_path, 'w:gz' if file_type == 'tar.gz' else 'w') as tf:
                            import tempfile
                            for i in range(3):
                                with tempfile.NamedTemporaryFile(mode='w', delete=False, encoding='utf-8') as tmp:
                                    tmp.write(fake.text())
                                    tmp.flush()
                                    tf.add(tmp.name, arcname=f'file_{i}.txt')
                                    os.unlink(tmp.name)
                    size = os.path.getsize(full_path)
                    success = True
                    
                elif file_type in ['sh', 'py', 'bash']:
                    # Scripts
                    if file_type == 'sh':
                        content = '#!/bin/bash\necho "Generated script"\n'
                    else:
                        content = '#!/usr/bin/env python3\nprint("Generated script")\n'
                    with open(full_path, 'w') as f:
                        f.write(content)
                    os.chmod(full_path, 0o755)
                    size = os.path.getsize(full_path)
                    success = True
                else:
                    # Generic
                    with open(full_path, 'wb') as f:
                        f.write(b'Generic file content\n')
                    size = os.path.getsize(full_path)
                    success = True
                
                if not success:
                    module.fail_json(msg=f"Error generating {full_name}")
                
                generated_files.append({
                    'path': full_path,
                    'name': full_name,
                    'file_type': file_type,
                    'index': global_index,
                    'size': size
                })
            except Exception as e:
                module.fail_json(msg=f"Error generating {full_name}: {str(e)}")
            
            global_index += 1
    
    # 4. Mark some files as deleted if deleted_ratio is specified
    files_to_delete = []
    if deleted_ratio > 0 and deleted_ratio <= 100:
        num_to_delete = int(len(generated_files) * (deleted_ratio / 100.0))
        if num_to_delete > 0:
            # Randomly select files to delete
            files_shuffled = generated_files.copy()
            random.shuffle(files_shuffled)
            files_to_delete = files_shuffled[:num_to_delete]
            
            # Delete the selected files using delete_file module
            for file_info in files_to_delete:
                try:
                    os.remove(file_info['path'])
                except Exception:
                    pass  # Ignore errors, file might not exist
    
    # 5. Summary
    module.exit_json(
        changed=True,
        generated_count=len(generated_files),
        files_deleted=len(files_to_delete),
        deleted_ratio_applied=deleted_ratio,
        applied_distribution=counts_per_type,
        structure_used=structure,
        base_directory=base_directory,
        faker_seed_used=faker_seed,
        files=generated_files[:10],  # First 10 as sample
        deleted_files=[f['path'] for f in files_to_delete][:10] if files_to_delete else [],
        total_files=len(generated_files)
    )


if __name__ == '__main__':
    main()
